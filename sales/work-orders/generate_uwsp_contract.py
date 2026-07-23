"""
UWSP Construction Contract generator (CWDB-2026-044).

Rebuilds the institutional (State-of-Wisconsin buyer) variant of the
CWDB-2026-044 deck-build contract as an editable .docx master, then exports a
signature-ready PDF via Microsoft Word (COM).

Why this exists: the original was hand-drafted in LibreOffice and only survived
as a flattened PDF. UWSP Procurement (Jesse Crain) will not sign with a binding-
arbitration clause (a WI state agency cannot agree to it by statute), and more
state-standard-terms edits may follow, so we need a real editable source.

Edits applied vs. the 2026-07-01 LibreOffice PDF (per legal-compliance-counsel
review, agent-memory/legal-compliance-counsel/state-agency-arbitration-rule.md):
  1. Section 13.2  Binding Arbitration  ->  Litigation (arbitration removed,
     statutory-rights savings language carried forward, explicit no-arbitration
     line added).
  2. Section 14.1  drop the trailing "for any proceeding not subject to
     arbitration" (stale once arbitration is gone).
  3. Exhibit A cover  "Home Improvement Contract"  ->  "Construction Contract"
     (matches the retitled body).
Everything else is preserved verbatim.

Usage:
    python generate_uwsp_contract.py [--effective "July 8, 2026"]
Outputs (next to this script):
    CWDB-2026-044-UWSP-Construction-Contract-v2.docx
    CWDB-2026-044-UWSP-Construction-Contract-v2.pdf   (if Word COM available)
"""

import argparse
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ORANGE = RGBColor(0xE5, 0x4C, 0x00)
SLATE = RGBColor(0x32, 0x34, 0x34)
GREY = RGBColor(0x64, 0x67, 0x60)

HERE = Path(__file__).resolve().parent
STEM = "CWDB-2026-044-UWSP-Construction-Contract-v2"


# ---------- low-level helpers ----------

def _set_font(run, *, size=None, bold=None, italic=None, color=None,
              name="Calibri"):
    run.font.name = name
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.font.bold = bold
    if italic is not None:
        run.font.italic = italic
    if color is not None:
        run.font.color.rgb = color


def _bottom_border(paragraph, color="E54C00", size="8"):
    p = paragraph._p
    pPr = p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), color)
    pbdr.append(bottom)
    pPr.append(pbdr)


def para(doc, text="", *, size=10.5, bold=False, italic=False, color=SLATE,
         align=None, space_after=6, space_before=0):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    if text:
        r = p.add_run(text)
        _set_font(r, size=size, bold=bold, italic=italic, color=color)
    return p


def section_heading(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    _set_font(r, size=13, bold=True, color=ORANGE)
    return p


def clause(doc, label, body, *, italic_body=False):
    """Run-in bold label ('13.1 Informal Resolution.') then body text."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    rl = p.add_run(label + " ")
    _set_font(rl, size=10.5, bold=True, color=SLATE)
    rb = p.add_run(body)
    _set_font(rb, size=10.5, italic=italic_body, color=SLATE)
    return p


def bullets(doc, items):
    for it in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(3)
        r = p.add_run(it)
        _set_font(r, size=10.5, color=SLATE)


def sig_block(doc, party_label, by_line):
    para(doc, party_label, bold=True, size=10.5, space_after=2)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run("_" * 52 + "        " + "_" * 22)
    _set_font(r, size=10.5, color=SLATE)
    q = doc.add_paragraph()
    q.paragraph_format.space_after = Pt(10)
    rq = q.add_run(by_line + " " * 40 + "Date")
    _set_font(rq, size=9, color=GREY)


def _shade_cell(cell, hexfill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hexfill)
    tcPr.append(shd)


def money_table(doc, header, rows, total_row):
    tbl = doc.add_table(rows=1, cols=len(header))
    tbl.style = "Table Grid"          # thin grid lines; header/total shaded manually
    tbl.autofit = True
    hdr = tbl.rows[0].cells
    for i, h in enumerate(header):
        hdr[i].text = ""
        rp = hdr[i].paragraphs[0].add_run(h)
        _set_font(rp, size=10, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))
        _shade_cell(hdr[i], "E54C00")            # CWDB orange header fill
    for row in rows:
        cells = tbl.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            rp = cells[i].paragraphs[0].add_run(val)
            _set_font(rp, size=10, color=SLATE)
    cells = tbl.add_row().cells
    for i, val in enumerate(total_row):
        cells[i].text = ""
        rp = cells[i].paragraphs[0].add_run(val)
        _set_font(rp, size=10, bold=True, color=SLATE)
        _shade_cell(cells[i], "FCE9DF")          # light-orange total row
    return tbl


# ---------- document ----------

def build(effective_label):
    doc = Document()
    doc.sections[0].top_margin = Inches(0.7)
    doc.sections[0].bottom_margin = Inches(0.7)
    doc.sections[0].left_margin = Inches(0.9)
    doc.sections[0].right_margin = Inches(0.9)
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)

    # Letterhead
    h = para(doc, "CENTRAL WISCONSIN DECK BUILDERS",
             size=15, bold=True, color=ORANGE,
             align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
    para(doc, "cwdeckbuilders.com  ·  (715) 544-7941  ·  info@cwdeckbuilders.com",
         size=9, color=GREY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=1)
    ruled = para(doc, "906 N 16th Ave, Wausau, WI 54401  ·  Entity ID C138564  ·  EIN 41-5355234",
                 size=9, color=GREY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=8)
    _bottom_border(ruled)

    para(doc, "CONSTRUCTION CONTRACT", size=16, bold=True, color=SLATE,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_before=8, space_after=2)
    para(doc, f"Job No. CWDB-2026-044  ·  Effective Date {effective_label}"
              "  ·  Ref. Estimate #2026-06-26-001",
         size=9.5, color=GREY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=10)

    # Preamble
    para(doc,
         "THIS CONSTRUCTION CONTRACT (this “Contract”) is entered into as of the "
         "Effective Date above, by and between Central Wisconsin Deck Builders, LLC, a "
         "Wisconsin limited liability company, 906 N 16th Ave, Wausau, WI 54401, phone "
         "(715) 544-7941, email info@cwdeckbuilders.com (“CWDB,” “Contractor,” or "
         "“we”), and the University of Wisconsin–Stevens Point, Purchasing "
         "Department, 2100 Main Street, Old Main Building, Room 041W, Stevens Point, WI "
         "54481-3897, email jcrain@uwsp.edu, attention: Jesse Crain, Director of "
         "Procurement (the “Owner” or “you”). CWDB and the Owner are each a "
         "“Party” and together the “Parties.” Salesperson / Representative: "
         "James Slogar.")

    # 1
    section_heading(doc, "Section 1. Parties and Project")
    clause(doc, "1.1 Project Property.",
           "CWDB will perform the work described in Section 2 at: 10186 County Road Mm, "
           "Amherst Junction, WI 54407 (the “Property”).")
    clause(doc, "1.2 Owner Authority.",
           "The Owner represents that the Owner owns the Property or is otherwise "
           "authorized to contract for the improvements described in this Contract.")
    clause(doc, "1.3 Contractor Identity.",
           "CWDB is the prime contractor for the Project. CWDB may perform the work "
           "itself or engage one or more subcontractors under a separate written "
           "subcontractor agreement. CWDB remains responsible to the Owner for the "
           "completed work regardless of who performs it.")

    # 2
    section_heading(doc, "Section 2. Scope of Work and Specifications")
    clause(doc, "2.1 Description of Work.",
           "CWDB will furnish the labor, materials, equipment, and supervision to "
           "perform the following work at the Property (the “Work”):")
    para(doc,
         "New deck build at the project address. The existing deck has already been "
         "removed by the Owner, so this scope covers new construction only. The job "
         "builds a new deck of approximately 128 sq ft of deck floor (16 ft x 8 ft) with "
         "32 linear feet of railing system and a staircase (1 run, 3 total treads). The "
         "new deck features Pressure-Treated Pine decking, Pressure-Treated Wood Rail "
         "railings, and kiln-dried pressure-treated lumber framing built to current "
         "Wisconsin residential code. Work includes complete site protection, "
         "code-compliant structure, and a final Owner walkthrough.")
    clause(doc, "2.2 Scope of Work.", "")
    bullets(doc, [
        "Site survey, layout, and Owner walkthrough of the staked footprint before any digging",
        "Move and protect Owner furniture, grill, plants, and adjacent landscaping",
        "Cover and protect siding, windows, doors, and any nearby paved surfaces",
        "Excavate and pour new code-compliant footings (12 in tube forms x 48 in deep for WI frost line) at code-required spacing",
        "Frame the new deck per WI residential code: kiln-dried pressure-treated lumber, ledger lag-bolted and flashed to house rim, code-compliant joist spacing on 16 in centers",
        "Install Pressure-Treated Pine decking boards with 3-inch stainless deck screws, maintaining manufacturer-specified gap spacing",
        "Install Pressure-Treated Wood Rail railing system (32 LF) with code-compliant baluster spacing and structural post anchoring",
        "Install staircase: 1 run, 3 treads, code-compliant rise and run",
        "Final touch-ups, complete cleanup, and removal of all construction debris",
        "Final Owner walkthrough and punch-list resolution",
    ])
    para(doc,
         "Scope boundary: The scope above assumes standard residential site access "
         "(driveway accessible to the delivery truck), no underground utilities within the "
         "deck footprint, level or near-level grade at the build location, and existing "
         "house siding/framing in sound condition at the ledger location. Any discoveries "
         "during excavation or framing (rot at the ledger location, inadequate or remaining "
         "old footings, undisclosed utilities) will be discussed before continuing work, "
         "and any change order will be priced separately and added in writing.",
         italic=True)
    clause(doc, "2.3 Detailed Specifications.",
           "The dimensions, configuration, materials, and finish specifications of the "
           "Work are as stated in Exhibit A (Estimate #2026-06-26-001), attached to and "
           "incorporated into this Contract.")
    clause(doc, "2.4 Included in the Price.", "")
    bullets(doc, [
        "All materials: kiln-dried pressure-treated lumber framing, Pressure-Treated Pine decking, Pressure-Treated Wood Rail railings, hardware (joist hangers, hurricane ties, lag bolts, fasteners), code-compliant flashing",
        "All new footings: concrete tube forms, concrete mix, rebar, and excavation",
        "All labor for framing, decking, railing, and stairs",
        "Permit and engineering allowance (Owner is named on the permit)",
        "Final job-site cleanup and haul-off of all construction debris",
        "10-year limited workmanship warranty on framing and structural connections",
        "Manufacturer warranties pass through on decking (Pressure-Treated Pine) and railing (Pressure-Treated Wood Rail)",
    ])
    clause(doc, "2.5 NOT Included in the Price.",
           "The following are not included and, if needed, will be handled by a signed "
           "change order under Section 5 or by others:")
    bullets(doc, [
        "Demolition or removal of any existing structure (the old deck has already been removed by the Owner)",
        "Landscape restoration beyond the immediate deck footprint (sod replacement, planting bed restoration, irrigation repair)",
        "Electrical, plumbing, gas, or hot tub utility hookups (separate trade required)",
        "Pergolas, screens, outdoor kitchens, or other built-ins unless added to scope in writing",
        "Structural modifications to the house itself (door replacement, header changes, siding alterations beyond the ledger flashing detail)",
        "Site grading, drainage modifications, or retaining wall work",
        "Any work required to remediate discoveries (rot, hidden utilities, code violations at ledger attachment, or removal of leftover old footings) quoted separately if found",
        "Painting or staining of new decking (wood decking can be quoted as an add-on for a first-year finish)",
        "Snow removal during winter installs",
    ])
    clause(doc, "2.6 Unforeseen Conditions.",
           "If CWDB discovers concealed or unforeseen conditions (for example rot, prior "
           "code violations, buried obstructions, or soil conditions requiring deeper "
           "footings) that materially affect cost or schedule, CWDB will stop affected "
           "work, notify the Owner, and proceed only under a signed change order.")

    # 3
    section_heading(doc, "Section 3. Total Price and Payment Schedule")
    clause(doc, "3.1 Total Price.",
           "The total price for the Work is $7,751 (the “Contract Price”). This is "
           "a fixed price subject only to signed change orders under Section 5.")
    clause(doc, "3.2 Payment Schedule.",
           "The Owner shall pay the Contract Price in the following installments:")
    money_table(doc,
                ["Milestone", "Amount", "When due"],
                [["Deposit at signing (30%)", "$2,325",
                  "Due upon execution of this Contract (and issuance of any required purchase order)"],
                 ["Final payment", "$5,426", "Upon final completion and walkthrough"]],
                ["Contract Price", "$7,751", ""])
    para(doc, "", space_after=2)
    clause(doc, "3.3 Deposit.",
           "The deposit covers initial materials and mobilization and is due upon "
           "execution of this Contract. CWDB may apply the deposit to job-specific "
           "materials and mobilization upon execution. The Owner is a "
           "governmental/institutional buyer contracting for non-consumer, commercial "
           "purposes; no consumer right of rescission applies to this Contract.")
    clause(doc, "3.4 Payment Method.",
           "Payments may be made by credit or debit card or digital payment (preferred), "
           "or check. Cash not accepted.")
    clause(doc, "3.5 Late Payment.",
           "Past-due amounts bear interest at one and one-half percent (1.5%) per month "
           "(18% per annum) or the maximum rate permitted by Wisconsin law, whichever is "
           "less.")
    clause(doc, "3.6 Lien Waivers.",
           "Upon the Owner’s request and upon payment, CWDB will provide a lien waiver "
           "for amounts paid, and will collect lien waivers from its subcontractors and "
           "suppliers, as described in Section 10.")

    # 4
    section_heading(doc, "Section 4. Schedule")
    clause(doc, "4.1 Start Date.",
           "CWDB estimates work will begin on or about August 2026, following execution "
           "of this Contract and issuance of any required purchase order.")
    clause(doc, "4.2 Substantial Completion.",
           "CWDB estimates the Work will reach substantial completion approximately "
           "7–10 working days after the start date, subject to weather, material "
           "availability, permit timing, change orders, and other causes beyond CWDB’s "
           "reasonable control. Work will be rescheduled if rain forecast exceeds 60% "
           "during exterior framing days. Winter installs may extend the duration by "
           "2–4 days due to weather windows. Concrete pours require ambient "
           "temperatures above 40 F for 48 hours after placement.")
    clause(doc, "4.3 Delays.",
           "CWDB is not responsible for delays caused by weather, acts of God, permit or "
           "inspection timing, material shortages, the Owner, or other causes beyond "
           "CWDB’s reasonable control. CWDB will give the Owner reasonable notice of any "
           "material delay and a revised estimated completion date.")

    # 5
    section_heading(doc, "Section 5. Change Orders")
    clause(doc, "5.1 Written Change Orders Only.",
           "Any change to the scope, materials, price, or schedule must be documented in "
           "a written change order signed by both Parties before the changed work "
           "proceeds. Change orders will reference this Contract’s Job No. "
           "(CWDB-2026-044/CO-1, CWDB-2026-044/CO-2, and so on).")
    clause(doc, "5.2 No Verbal Changes.",
           "CWDB is not obligated to perform, and the Owner is not obligated to pay for, "
           "any extra or changed work that is not covered by a signed change order. Verbal "
           "requests do not bind either Party.")
    clause(doc, "5.3 Effect on Price and Schedule.",
           "Each change order will state the change in scope, the increase or decrease in "
           "the Contract Price, and any change to the schedule.")

    # 6
    section_heading(doc, "Section 6. Permits and Code Compliance")
    clause(doc, "6.1 Permits.",
           "CWDB will obtain and pay for the building permit(s) required for the Work and "
           "the cost is included in the Contract Price. The responsible Party will obtain "
           "all permits required by Marathon County and the applicable municipality before "
           "starting permitted work.")
    clause(doc, "6.2 Code Compliance.",
           "CWDB will perform the Work in accordance with the Wisconsin Uniform Dwelling "
           "Code (Wis. Admin. Code SPS 320 through 325), including deck-specific "
           "structural requirements for footings, ledger attachment, guardrail height, and "
           "load rating, and in accordance with applicable local building codes and "
           "ordinances.")
    clause(doc, "6.3 Inspections.",
           "CWDB will coordinate required inspections. The Owner will provide reasonable "
           "access to the Property for inspections and for the Work.")

    # 7
    section_heading(doc, "Section 7. Warranty")
    clause(doc, "7.1 Workmanship Warranty.",
           "CWDB warrants that the Work will be free from defects in workmanship for a "
           "period of ten (10) years from the date of substantial completion. If a covered "
           "workmanship defect appears during the warranty period, CWDB will, at its "
           "option and at no charge to the Owner, repair or correct the defect, provided "
           "the Owner gives written notice of the defect within the warranty period.")
    clause(doc, "7.2 Manufacturer Warranties Pass Through.",
           "Materials and products installed by CWDB may carry separate manufacturer "
           "warranties. CWDB passes through to the Owner all such manufacturer warranties "
           "to the extent transferable. CWDB does not separately warrant the materials "
           "beyond the workmanship warranty in Section 7.1 and the express manufacturer "
           "warranties.")
    clause(doc, "7.3 Exclusions.",
           "The warranty does not cover: normal weathering, fading, or wear; damage from "
           "misuse, neglect, or failure to maintain (including failure to re-stain or seal "
           "on a normal maintenance schedule); movement, settling, or splitting of natural "
           "wood that is consistent with industry norms; damage from storms, floods, or "
           "other events beyond CWDB’s control; or work, alteration, or repair performed "
           "by anyone other than CWDB after completion.")
    clause(doc, "7.4 Maintenance.",
           "The Owner is responsible for routine maintenance, including periodic cleaning "
           "and re-staining or re-sealing as recommended for the chosen product.")

    # 8
    section_heading(doc, "Section 8. Cleanup and Debris")
    clause(doc, "8.1 Cleanup.",
           "CWDB will keep the work area reasonably clean during the Work and, upon "
           "completion, will remove its tools, equipment, surplus materials, and "
           "construction debris and leave the work area in broom-clean condition.")
    clause(doc, "8.2 Disposal.",
           "CWDB will lawfully dispose of construction debris generated by the Work. "
           "Disposal of pre-existing materials or hazardous materials not generated by the "
           "Work is not included unless stated in Section 2.")

    # 9
    section_heading(doc, "Section 9. Insurance")
    clause(doc, "9.1 CWDB Insurance.",
           "CWDB represents that it maintains commercial general liability insurance "
           "covering its operations. CWDB will provide a certificate of insurance to the "
           "Owner upon request.")
    clause(doc, "9.2 Subcontractor Insurance.",
           "Any subcontractor CWDB engages for the Work is required, under a separate "
           "subcontractor agreement, to carry its own commercial general liability "
           "insurance and workers’ compensation coverage (or a valid Wisconsin "
           "exemption) and to name CWDB as an additional insured before starting work.")

    # 10
    section_heading(doc, "Section 10. Lien Waivers")
    clause(doc, "10.1 Lien Waivers.",
           "As the Work progresses and as payments are made, CWDB will, upon request, "
           "provide lien waivers and will collect lien waivers from its subcontractors and "
           "suppliers to protect the Owner’s clear title to the extent applicable to the "
           "Property.")
    clause(doc, "10.2 Security Interest.",
           "CWDB takes no security interest, mortgage, or lien on the Property or on any "
           "of the Owner’s property as part of this Contract, apart from any statutory "
           "construction-lien rights that arise by operation of law for unpaid labor and "
           "materials and only to the extent such rights apply to the Property.")

    # 11
    section_heading(doc, "Section 11. Warranties, Disclaimers, and Limitation of Liability")
    clause(doc, "11.1 Exclusive Warranty.",
           "The express warranty in Section 7 is the only warranty CWDB provides. To the "
           "fullest extent permitted by Wisconsin law, and except for the express warranty "
           "in Section 7 and any nondisclaimable warranties under Wisconsin law, CWDB "
           "disclaims all other warranties, whether express or implied, including any "
           "implied warranty of merchantability or fitness for a particular purpose, beyond "
           "the express warranty stated in this Contract.")
    clause(doc, "11.2 Limitation of Liability.",
           "TO THE FULLEST EXTENT PERMITTED BY WISCONSIN LAW, CWDB SHALL NOT BE LIABLE TO "
           "THE OWNER FOR ANY INDIRECT, INCIDENTAL, SPECIAL, CONSEQUENTIAL, PUNITIVE, OR "
           "EXEMPLARY DAMAGES, INCLUDING DAMAGES FOR LOSS OF USE, INCONVENIENCE, OR LOST "
           "PROFITS, ARISING OUT OF OR RELATING TO THIS CONTRACT OR THE WORK. CWDB’s "
           "total aggregate liability under this Contract shall not exceed the Contract "
           "Price actually paid by the Owner. Nothing in this Section limits any liability "
           "that cannot be limited under Wisconsin law, including liability for personal "
           "injury caused by CWDB’s negligence.")

    # 12
    section_heading(doc, "Section 12. Indemnification")
    clause(doc, "12.1 CWDB Responsibility for Its Work.",
           "CWDB will be responsible for property damage or bodily injury to the extent "
           "caused by the negligent acts or omissions of CWDB or its subcontractors in "
           "performing the Work, subject to the limitations in Section 11.")
    clause(doc, "12.2 Owner Responsibility.",
           "The Owner will hold CWDB harmless from claims arising out of conditions at the "
           "Property that the Owner failed to disclose, the Owner’s own acts or "
           "omissions, or the acts of persons the Owner brings onto the Property who are "
           "not under CWDB’s control, to the extent permitted by Wisconsin law "
           "applicable to the Owner.")

    # 13  ---- EDITED SECTION ----
    section_heading(doc, "Section 13. Dispute Resolution")
    clause(doc, "13.1 Informal Resolution.",
           "Before starting any formal proceeding, the Party raising a dispute will give "
           "the other written notice, and the Parties will try in good faith to resolve "
           "the dispute through informal negotiation for thirty (30) calendar days.")
    # 13.2: arbitration removed; litigation clause per legal-compliance-counsel.
    clause(doc, "13.2 Litigation.",
           "If the dispute is not resolved within that thirty (30) calendar day period, "
           "either Party may pursue any available legal remedy in a court of competent "
           "jurisdiction, subject to the governing law and venue provisions of Section "
           "14.1. The Parties agree that no dispute arising under this Contract is subject "
           "to binding arbitration. Nothing in this Section limits or waives either "
           "Party’s rights or remedies under applicable Wisconsin law, including any "
           "claim procedures applicable to the State of Wisconsin, the University of "
           "Wisconsin System, and their agencies.")
    clause(doc, "13.3 Equitable Relief.",
           "Either Party may seek injunctive relief in a court of competent jurisdiction "
           "to prevent irreparable harm without first completing the steps in this "
           "Section.")

    # 14
    section_heading(doc, "Section 14. General Provisions")
    clause(doc, "14.1 Governing Law; Venue.",
           "This Contract is governed by the laws of the State of Wisconsin. The Parties "
           "consent to the exclusive jurisdiction and venue of the state and federal "
           "courts located in Wisconsin.")  # arbitration tail removed
    clause(doc, "14.2 Entire Agreement.",
           "This Contract, together with Exhibit A (Estimate #2026-06-26-001) and any "
           "signed change orders, is the entire agreement between the Parties and "
           "supersedes all prior discussions, estimates, and representations. If the Owner "
           "issues a purchase order for the Work, the Owner’s standard purchase-order "
           "terms and conditions, to the extent the Parties agree in writing, are "
           "incorporated by reference and, in the event of a conflict, govern to the "
           "extent required by law.")
    clause(doc, "14.3 No Reliance.",
           "CWDB has not made any false, deceptive, or misleading representation regarding "
           "the Work, and the Owner is not relying on any promise, representation, or "
           "warranty not written in this Contract.")
    clause(doc, "14.4 Amendments; Severability; No Waiver.",
           "This Contract may be amended only by a written instrument or signed change "
           "order executed by both Parties. If any provision is held invalid or "
           "unenforceable, it will be modified to the minimum extent necessary to make it "
           "enforceable, and the remaining provisions will remain in effect. A Party’s "
           "failure to enforce a provision is not a waiver of the right to enforce it "
           "later.")
    clause(doc, "14.5 Assignment.",
           "The Owner may not assign this Contract without CWDB’s written consent. CWDB "
           "may use subcontractors as provided in Section 1.3 but remains responsible for "
           "the Work.")
    clause(doc, "14.6 Notices.",
           "Notices must be in writing and delivered personally, by email with "
           "confirmation, or by certified mail to the addresses in this Contract.")
    clause(doc, "14.7 Counterparts and Electronic Signatures.",
           "This Contract may be signed in counterparts and by electronic signature, each "
           "of which is an original.")

    # Signatures
    section_heading(doc, "Signatures")
    para(doc,
         "IN WITNESS WHEREOF, the Parties have executed this Contract as of the Effective "
         "Date. By signing, the Owner acknowledges reading and agreeing to this Contract, "
         "and that the person signing is authorized to bind the Owner.", space_after=10)
    sig_block(doc, "Contractor: Central Wisconsin Deck Builders, LLC",
              "By: James Slogar, Member")
    sig_block(doc, "Owner: University of Wisconsin–Stevens Point",
              "By: Jesse Crain, Director of Procurement")
    para(doc, "Printed name: Jesse Crain", size=10.5, space_after=1)
    para(doc, "Title: Director of Procurement", size=10.5, space_after=1)

    # ---------- Exhibit A ----------
    doc.add_page_break()
    para(doc, "EXHIBIT A", size=15, bold=True, color=SLATE,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
    # label fixed: "Home Improvement Contract" -> "Construction Contract"
    para(doc,
         "Estimate #2026-06-26-001, incorporated into and made part of Construction "
         "Contract Job No. CWDB-2026-044",
         size=10, color=GREY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=10)
    ruled2 = para(doc,
                  "PROJECT ESTIMATE  ·  #2026-06-26-001  ·  Issued June 26, 2026  "
                  "·  Valid 14 days",
                  size=10, bold=True, color=SLATE,
                  align=WD_ALIGN_PARAGRAPH.CENTER, space_after=8)
    _bottom_border(ruled2)

    para(doc, "PREPARED FOR", size=10, bold=True, color=ORANGE, space_after=1)
    for ln in ["University of Wisconsin–Stevens Point",
               "Attn: Jesse Crain, Director of Procurement",
               "Purchasing Department",
               "2100 Main Street, Old Main Building, Room 041W",
               "Stevens Point, WI 54481-3897",
               "jcrain@uwsp.edu"]:
        para(doc, ln, size=10, space_after=0)
    para(doc, "ESTIMATOR", size=10, bold=True, color=ORANGE, space_before=6, space_after=1)
    for ln in ["James Slogar, Owner",
               "Central Wisconsin Deck Builders, LLC",
               "(715) 544-7941",
               "info@cwdeckbuilders.com"]:
        para(doc, ln, size=10, space_after=0)

    section_heading(doc, "Project Overview")
    para(doc,
         "New deck build at the project address. The existing deck has already been "
         "removed by the Owner, so this scope covers new construction only. The job builds "
         "a new deck of approximately 128 sq ft of deck floor (16 ft x 8 ft) with 32 "
         "linear feet of railing system and a staircase (1 run, 3 total treads). The new "
         "deck features Pressure-Treated Pine decking, Pressure-Treated Wood Rail "
         "railings, and kiln-dried pressure-treated lumber framing built to current "
         "Wisconsin residential code. Work includes complete site protection, "
         "code-compliant structure, and a final Owner walkthrough.")

    section_heading(doc, "Itemized Pricing")
    money_table(doc,
                ["Item", "Price"],
                [["Deck construction including KDAT Pressure-Treated framing, "
                  "Pressure-Treated Pine decking (128 sq ft)", "$3,368"],
                 ["Pressure-Treated Wood Rail railing system (32 LF)", "$784"],
                 ["Staircase: 1 run, 3 treads", "$1,608"],
                 ["Permits, site protection, mobilization, and final cleanup", "$1,991"]],
                ["TOTAL FIXED PRICE", "$7,751"])

    section_heading(doc, "Payment Terms")
    money_table(doc,
                ["Stage", "Amount"],
                [["Deposit due on signed acceptance (30%)", "$2,325"],
                 ["Balance due on completion and walkthrough (70%)", "$5,426"]],
                ["Total", "$7,751"])
    para(doc, "Accepted forms of payment: credit or debit card or digital payment "
              "(preferred), or check. Cash not accepted.",
         size=9.5, italic=True, color=GREY, space_before=3)

    section_heading(doc, "Acceptance")
    para(doc,
         "Acceptance of this estimate is subject to execution of CWDB’s signed "
         "Construction Contract, which will carry the deposit terms stated above. This "
         "estimate is not itself a binding contract. Estimate is valid for 14 days from the "
         "date issued (June 26, 2026).", space_after=12)
    sig_block(doc, "Owner: University of Wisconsin–Stevens Point",
              "By: Jesse Crain, Director of Procurement")
    sig_block(doc, "Central Wisconsin Deck Builders, LLC",
              "By: James Slogar, Owner")

    return doc


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("--effective", default="_____________, 2026",
                    help="Effective Date shown in the header (blank to date at signing)")
    args = ap.parse_args()

    docx_path = HERE / f"{STEM}.docx"
    doc = build(args.effective)
    doc.save(str(docx_path))
    print("WROTE", docx_path)


if __name__ == "__main__":
    main()
